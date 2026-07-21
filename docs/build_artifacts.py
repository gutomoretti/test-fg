from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"

def font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()

def diagram():
    image = Image.new("RGB", (1500, 920), "white")
    draw = ImageDraw.Draw(image)
    title = font(32, True)
    text = font(22)
    small = font(19)
    blue = "#1F4D78"
    line = "#4F6272"

    def box(x, y, w, h, name, lines, fill="#F4F6F9"):
        draw.rounded_rectangle((x, y, x+w, y+h), radius=12, fill=fill, outline=blue, width=3)
        draw.rectangle((x, y, x+w, y+54), fill="#E8EEF5", outline=blue, width=3)
        draw.text((x+18, y+12), name, font=font(24, True), fill="#0B2545")
        yy = y + 72
        for item in lines:
            draw.text((x+18, yy), item, font=small, fill="#1D2733")
            yy += 32

    draw.text((52, 28), "Diagrama de classes - Serviço de ordenação", font=title, fill="#0B2545")
    box(70, 160, 360, 300, "Book", ["+ Title: string", "+ AuthorName: string", "+ EditionYear: int"], "#F8FBFD")
    box(555, 145, 390, 350, "BooksOrderer", ["<<interface>>", "+ Order(books): IReadOnlyList<Book>"], "#EEF5FA")
    box(1090, 120, 350, 410, "ConfiguredBooksOrderer", ["- comparers", "+ ConfiguredBooksOrderer", "+ Order(books)", "- CompareBooks()", "- CreateComparer()"], "#F8FBFD")
    box(80, 610, 390, 210, "OrderConfiguration", ["+ Criteria: List<OrderCriterion>", "+ FromJsonFile(path)"], "#F8FBFD")
    box(610, 620, 350, 190, "OrderCriterion", ["+ Attribute: string", "+ Direction: SortDirection"], "#F8FBFD")
    box(1110, 650, 310, 150, "OrdenacaoException", ["Exception de domínio"], "#FFF8EA")

    def arrow(points, label=None, label_pos=None):
        draw.line(points, fill=line, width=4)
        x1, y1 = points[-2]; x2, y2 = points[-1]
        import math
        angle = math.atan2(y2-y1, x2-x1)
        for delta in (2.6, -2.6):
            draw.line([(x2, y2), (x2-18*math.cos(angle+delta), y2-18*math.sin(angle+delta))], fill=line, width=4)
        if label:
            draw.text(label_pos, label, font=small, fill="#34495E")

    arrow([(430, 300), (555, 300)], "recebe Book", (445, 265))
    arrow([(945, 300), (1090, 300)], "implementa", (965, 265))
    arrow([(750, 495), (750, 620)], "usa", (775, 535))
    arrow([(1090, 475), (960, 700)], "configura", (990, 560))
    arrow([(1260, 530), (1260, 650)], "lança", (1280, 575))
    image.save(DOCS / "uml-ordenacao-livros.png")

def document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor(31, 39, 48)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]; style.font.name = "Calibri"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Avaliação Técnica FGV\nServiço de Ordenação de Livros"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor.from_string("0B2545")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento de projeto e implementação"); r.italic = True; r.font.size = Pt(12)
    doc.add_picture(str(DOCS / "uml-ordenacao-livros.png"), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figura 1 - Estrutura principal da solução."); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph("A solução implementa um serviço .NET para ordenar livros por um ou mais atributos, com direção ascendente ou descendente definida por configuração externa.")
    doc.add_heading("2. Arquitetura", level=1)
    doc.add_paragraph("A biblioteca Fgv.Ordenacao concentra o domínio e a regra de ordenação. O projeto Fgv.Ordenacao.Api fornece o host ASP.NET Core e o endpoint HTTP. O projeto Fgv.Ordenacao.Tests contém testes unitários e testes de integração.")
    doc.add_heading("3. Decisões de projeto", level=1)
    doc.add_heading("Contrato público pequeno", level=2)
    doc.add_paragraph("BooksOrderer expõe somente Order(IEnumerable<Book>). A interface não conhece HTTP, JSON, Swagger ou a forma como os critérios são armazenados, mantendo a biblioteca reutilizável e desacoplada.")
    doc.add_heading("Ordenação composta", level=2)
    doc.add_paragraph("ConfiguredBooksOrderer cria um comparador para cada critério configurado e os avalia na ordem declarada. Quando um critério encontra diferença, seu resultado é retornado; caso contrário, o próximo critério é avaliado.")
    doc.add_heading("Configuração e validação", level=2)
    doc.add_paragraph("O arquivo ordering.json define os atributos e direções. Entradas nulas, coleções vazias, configuração sem critérios e atributos desconhecidos são tratados por OrdenacaoException e retornam 400 Bad Request na API.")
    doc.add_heading("4. API", level=1)
    doc.add_paragraph("POST /books/order recebe livros em JSON e retorna a coleção ordenada. Em desenvolvimento, a documentação interativa está disponível em /swagger.")
    doc.add_paragraph("Exemplo de corpo: { books: [ { title, authorName, editionYear } ] }")
    doc.add_heading("5. Testes e execução", level=1)
    doc.add_paragraph("A suíte cobre os critérios fornecidos pela avaliação, entradas inválidas, leitura da configuração e o fluxo HTTP. Execute com:")
    p = doc.add_paragraph(); p.style = "No Spacing"; p.add_run("dotnet test Fgv.Ordenacao.sln --configuration Release").font.name = "Consolas"
    doc.add_paragraph("Para iniciar a API:")
    p = doc.add_paragraph(); p.style = "No Spacing"; p.add_run("dotnet run --project src/Fgv.Ordenacao.Api --configuration Release").font.name = "Consolas"
    doc.add_heading("6. Entregáveis", level=1)
    doc.add_paragraph("O pacote contém o código-fonte completo, configuração externa, testes automatizados, este documento e o diagrama UML em formato PNG.")
    doc.save(DOCS / "documentacao-tecnica.docx")

if __name__ == "__main__":
    diagram(); document()
